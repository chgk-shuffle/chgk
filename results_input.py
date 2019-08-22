import sys
import tour
import sqlite3
from random import shuffle
from docx import Document
from docx.shared import Inches
import psycopg2
import sqlite3
import urllib.parse as urlparse
import os

from ui.results_input import *
from PyQt5 import QtCore, QtGui, QtWidgets, Qt


class ButtonBlock(QtWidgets.QWidget):
    def __init__(self, *args):
        tour = args[0]
        ques = args[1]
        super(QtWidgets.QWidget, self).__init__()
        global nTeam
        self.grid = QtWidgets.QGridLayout()
        for g in range(nTeam):
            pushButton = QtWidgets.QPushButton(self)
            pushButton.setText("Команда " + str(g + 1))
            pushButton.setCheckable(1)
            font = pushButton.font()
            font.setPointSize(25)
            pushButton.setFont(font)
            pushButton.tour = tour
            pushButton.ques = ques
            pushButton.clicked.connect(self.make_calluser(g))
            self.grid.addWidget(pushButton, g // 4, g % 4)

        self.setLayout(self.grid)

    def make_calluser(self, curT):
        def calluser():
            sender = self.sender()
            if sender.isChecked():
                point1 = 1
            else:
                point1 = -1
            cur.execute(f'''UPDATE "Team" SET score = score + {point1}
                            WHERE id_table = {curT}''')
            cur.execute(f'''UPDATE "User" SET score = score + {point1}
                            WHERE id_user IN (SELECT id_user FROM "Team"
                                                WHERE id_table = {curT}
                                                AND id_round = {sender.tour}
                                                )''')
            cur.execute(f'''UPDATE "Question" SET point = {1 if point1 == 1 else 0}
                            WHERE id_round = {sender.tour} AND question_number = {sender.ques}
                            AND id_table = {curT}''')
            con.commit()

        return calluser

    def get(self):
        return self


class MyWin(QtWidgets.QMainWindow):
    def __init__(self, parent=None):
        global con, cur
        con = psycopg2.connect("dbname='def6ihrm3usufr' "
                               "user='vwpffneyqlsshw'"
                               "host='ec2-54-235-86-101.compute-1.amazonaws.com' "
                               "password='1636d214d3260f0e48bc3ea4cbbd3912da56c3d6bad56c0b4dc05f360a2e4acc'"
                               "port='5432'"
                               )
        cur = con.cursor()
        QtWidgets.QWidget.__init__(self, parent, QtCore.Qt.Window)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.btnPrint1.clicked.connect(self.PrintRealList1)
        self.ui.btnPrint2.clicked.connect(self.PrintRealList2)
        self.curRound = 0
        self.curQ = 0
        global TabFinallRes
        global ResInTourTab

        cur.execute('''SELECT name, id_user FROM "User"''')
        self.listUsers = cur.fetchall()

        ResInTourTab = []
        count = 0
        cur.execute('''SELECT count(*) FROM "User"''')
        a = cur.fetchone()[0]
        cur.execute('''SELECT id FROM "Team_size"''')
        b = cur.fetchall()[0][0]
        self.teamsS = foo(a, b)

        quantityTeam = len(self.teamsS)
        self.tours = []
        cur.execute('''SELECT * FROM "Round"''')
        Round_list = cur.fetchall()
        print(Round_list)
        for i in Round_list:
            self.tours.append(tour.Tour(quantityTeam, i[3] - i[2] + 1, i[0], i[1]))
        self.finall_distrib = []
        self.countResTour = [0, 0]
        for i in self.tours:
            global nTeam
            nTeam = len(i.teamScore)
            tab = QtWidgets.QWidget()
            verticalLayout = QtWidgets.QVBoxLayout(tab)
            verticalLayout.setSizeConstraint(QtWidgets.QLayout.SetDefaultConstraint)
            tabWidget = QtWidgets.QTabWidget(tab)
            self.countResTour.append(0)
            DistribTab = QtWidgets.QTableWidget(nTeam, 1)
            DistribTab.setHorizontalHeaderLabels(["Участники"])
            index = 0
            shuffle(self.listUsers)
            shuffle(self.teamsS)
            lcount = 0

            list_string = []
            for j in self.teamsS:
                string = '\n'
                for v in range(index, index + j):
                    string += self.listUsers[v][0] + '\n'
                    # cur.execute(f'''INSERT INTO "Team"
                    #                 VALUES ({lcount}, {listUsers[v][1]}, {i.id}, 0)''')
                con.commit()
                list_string.append(string)
                DistribTab.setItem(lcount, 0, QtWidgets.QTableWidgetItem(string))
                index += j
                lcount += 1
            self.finall_distrib.append(list_string)

            DistribTab.resizeColumnsToContents()
            DistribTab.resizeRowsToContents()
            tabWidget.addTab(DistribTab, "Распределение")

            if i.type == 1:
                for j in range(i.countQ):
                    tabQuestion = QtWidgets.QWidget()
                    VerL = QtWidgets.QVBoxLayout(tabQuestion)
                    bt = ButtonBlock((i.id), j)
                    VerL.addWidget(bt)

                    tabWidget.addTab(tabQuestion, str(j + 1) + " вопрос")
                tabWidget.currentChanged.connect(self.onChangeResTourTabCHGK)
                ResInTourTab.append(QtWidgets.QTableWidget(nTeam, i.countQ + 2))
                ResInTourTab[i.id].setHorizontalHeaderLabels(
                    ['Команды'] + ['Вопрос ' + str(g + 1) for g in range(i.countQ)] + ['Сумма'])
            else:
                tabWidget.currentChanged.connect(self.onChangeResTourTabQVIZ)
                ResInTourTab.append(QtWidgets.QTableWidget(nTeam, 2))
                ResInTourTab[i.id].setHorizontalHeaderLabels(
                    ['Команды'] + ['Сумма'])
                ResInTourTab[i.id].cellChanged.connect(self.OnChangedItemTable)

            ResInTourTab[i.id].setItem(0, 0, QtWidgets.QTableWidgetItem('команда 13'))
            ResInTourTab[i.id].resizeColumnsToContents()
            tabWidget.addTab(ResInTourTab[i.id], "Итоги тура")
            verticalLayout.addWidget(tabWidget)
            verticalLayout.addWidget(tabWidget)
            self.ui.tabWidget.addTab(tab, str(count + 1) + " тур")
            count += 1

        self.ui.tabWidget.currentChanged.connect(self.onChangeTourTab)
        TabFinallRes = QtWidgets.QTableWidget(len(self.listUsers), 2)
        TabFinallRes.setHorizontalHeaderLabels(['Фамилия и Имя', 'Кол-во очков'])

        TabFinallRes.horizontalHeader().sectionClicked.connect(self.onChangeHeadResTab)
        TabFinallRes.horizontalHeader().sectionDoubleClicked.connect(self.onChangeHeadResTab2)
        self.LoadRes(0)
        self.ui.tabWidget.addTab(TabFinallRes, "Итоги")

        self.ui.tabWidget.addTab(QtWidgets.QTabWidget(), 'Справка')

    def onChangeTourTab(self, i):
        i -= 1
        self.countResTour[self.curRound] = 0
        self.curRound = i
        if i == len(self.tours):
            self.LoadRes(0)

    def onChangeResTourTabCHGK(self, i):
        if (i > 0) and (i < self.tours[self.curRound].countQ + 1):
            self.curQ = i
        if i == self.tours[self.curRound].countQ + 1:
            self.LoadResTourCHGK(self.tours[self.curRound].countQ)

    def onChangeResTourTabQVIZ(self, i):
        if i == 1:
            self.LoadResTourQVIZ()

    def onChangeHeadResTab(self, logicalIndex):
        if logicalIndex == 0:
            self.LoadRes(-1)
        else:
            self.LoadRes(1)

    def onChangeHeadResTab2(self, logicalIndex):
        if logicalIndex == 0:
            self.LoadRes(-2)
        else:
            self.LoadRes(2)

    def LoadResTourCHGK(self, nQ):
        cur.execute(f'''
                                            SELECT id_table, point, question_number FROM "Question"
                                            WHERE id_round = {self.curRound}
                                            ORDER BY id_table
                                            ''')
        ResInTour = cur.fetchall()
        ResInTourInTable = [[0 for g in range(nQ + 2)] for _ in range(nTeam)]
        for j in range(nTeam):
            ResInTourInTable[j][0] = j + 1
        for j in range(len(ResInTour)):
            ResInTourInTable[ResInTour[j][0]][ResInTour[j][2] + 1] = ResInTour[j][1]
            ResInTourInTable[ResInTour[j][0]][nQ + 1] = sum(ResInTourInTable[ResInTour[j][0]][1:nQ + 1])
        ResInTourInTable.sort(key=lambda team: team[nQ + 1], reverse=True)
        for j in range(len(ResInTourInTable)):
            for k in range(len(ResInTourInTable[j])):
                item = QtWidgets.QTableWidgetItem(
                    'команда ' + str(ResInTourInTable[j][k]) if (k == 0) else str(ResInTourInTable[j][k]) if (
                            k == nQ + 1) else '+' if ResInTourInTable[j][k] == 1 else '-')
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                ResInTourTab[self.curRound].setItem(j, k, item)

    def LoadResTourQVIZ(self):
        cur.execute(f'''
                                                            SELECT id_table, score FROM "Team"
                                                            WHERE id_round = {self.curRound}
                                                            ORDER BY id_table
                                                            ''')
        ResInTour = cur.fetchall()
        last = -1
        Res = []
        for i, j in ResInTour:
            if i != last:
                Res.append([i, j])
            last = i
        Res.sort(key=lambda e: e[1], reverse=True)
        for j in range(len(Res)):
            item = QtWidgets.QTableWidgetItem('команда ' + str(Res[j][0] + 1))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            ResInTourTab[self.curRound].setItem(j, 0, item)
            item = QtWidgets.QTableWidgetItem(str(Res[j][1]))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            ResInTourTab[self.curRound].setItem(j, 1, item)

    def LoadRes(self, type):
        if type == 0:
            sort = "score DESC, name"
        elif type == -1:
            sort = "name"
        elif type == 1:
            sort = "score DESC"
        elif type == -2:
            sort = "name DESC"
        else:
            sort = "score"
        cur.execute(f'''
                                SELECT name, score FROM "User"
                                ORDER BY {sort}
                                ''')
        FinallRes = cur.fetchall()
        for i in range(len(FinallRes)):
            for j in range(len(FinallRes[i])):
                item = QtWidgets.QTableWidgetItem(str(FinallRes[i][j]))
                item.setTextAlignment(QtCore.Qt.AlignCenter)
                TabFinallRes.setItem(i, j, item)
        TabFinallRes.resizeColumnsToContents()

    def PrintRealList1(self):
        document = Document()
        head = document.add_heading('Распределение участников', 0)
        head.alignment = 1
        table = document.add_table(rows=0, cols=3)
        table.style = 'TableGrid'
        for i in range(len(self.teamsS)):
            hdr_cells = table.add_row().cells
            hdr_cells[0].text = str(i + 1) + ' Стол'
            hdr_cells[1].text = 'Тур'
            hdr_cells[2].text = 'Участники'
            for j in range(len(self.tours)):
                row_cells_tour = table.add_row().cells

                row_cells_tour[1].text = str(j + 1)
                string = self.finall_distrib[j][i]
                row_cells_tour[2].text = string[1:].replace('\n', '; ')

        document.add_page_break()
        document.save('Распределение_по_турам.docx')

    def PrintRealList2(self):
        document = Document()
        cur.execute('SELECT id_user, name FROM "User"')
        names = dict(cur.fetchall())
        for id_user, name in names.items():
            document.add_heading(f'{name}', level=1)
            a = f'SELECT  id_user, id_table, id_round FROM "Team" WHERE id_user = {id_user} ORDER BY id_round'
            cur.execute(a)
            a = cur.fetchall()
            print(a)
            paragraph = document.add_paragraph()
            for i in a:
                lst = list(i)
                tabel = lst[1]
                round = lst[2]
                p = paragraph.add_run(f'Раунд {round}, стол {tabel} -----> ')
            p = document.add_paragraph(
                '_______________________________________________________________________________________________________')
        document.save('Маршрутный_лист.docx')

    def OnChangedItemTable(self, x, y):
        if self.countResTour[self.curRound] < len(self.teamsS):
            if y == 1:
                self.countResTour[self.curRound] += 1
        else:
            self.countResTour[self.curRound] = 0
            if y == 1:
                new_val = int(ResInTourTab[self.curRound].item(x, y).text())
                cur.execute(f'''SELECT score FROM "Team"
                                                                                WHERE id_table = {x}
                                                                                AND id_round = {self.curRound}
                                                                                ''')
                prev = cur.fetchone()[0]
                cur.execute(f'''UPDATE "Team" SET score = {new_val}
                                            WHERE id_table = {x} AND id_round = {self.curRound}''')
                cur.execute(f'''UPDATE "User" SET score = score + {new_val - prev}
                                            WHERE id_user IN (SELECT id_user FROM "Team"
                                                                WHERE id_table = {x}
                                                                AND id_round = {self.curRound}
                                                                )''')
                con.commit()


def foo(quantityUser, sizeTeam):
    teamsS = []
    a = [-1 for _ in range(0, quantityUser + 1)]
    a[0] = 0
    while a[quantityUser] == -1:
        for v in range(0, quantityUser - sizeTeam + 1):
            if a[v] != -1:
                a[v + sizeTeam] = max(sizeTeam, a[v + sizeTeam])
        sizeTeam -= 1
    index = quantityUser
    CurSizeTeam = a[index]
    while CurSizeTeam != 0:
        teamsS.append(CurSizeTeam)
        index -= CurSizeTeam
        CurSizeTeam = a[index]

    return teamsS


if __name__ == "__main__":
    global con, cur
    try:
        con = psycopg2.connect("dbname='def6ihrm3usufr' "
                               "user='vwpffneyqlsshw'"
                               "host='ec2-54-235-86-101.compute-1.amazonaws.com' "
                               "password='1636d214d3260f0e48bc3ea4cbbd3912da56c3d6bad56c0b4dc05f360a2e4acc'"
                               "port='5432'"
                               )
        cur = con.cursor()
    except:
        print("I am unable to connect to the database")

    app = QtWidgets.QApplication(sys.argv)
    app.setStyle('Fusion')
    myapp = MyWin()
    myapp.show()

    sys.exit(app.exec_())
