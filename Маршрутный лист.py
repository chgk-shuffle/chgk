import sqlite3
from docx import Document

def main():
    document = Document()
    cur.execute("SELECT id_user, name FROM USER ")
    names = dict(cur.fetchall())
    for id_user, name in names.items():
        document.add_heading(f'{name}', level=1)
        a = f"SELECT  id_user, id_table, id_round FROM Team WHERE id_user = {id_user} ORDER BY id_round"
        cur.execute(a)
        a = cur.fetchall()
        print(a)
        for i in a:
            lst = list(i)
            tabel = lst[1]
            round = lst[2]
            p = document.add_paragraph(f'Раунд {round}, стол {tabel}')
        p = document.add_paragraph('_______________________________________________________________________________________________________')
    document.save('Маршрутный лист.docx')

if __name__ == "__main__":
    global con, cur
    con = sqlite3.connect('db')
    cur = con.cursor()
    con.commit()
    main()
